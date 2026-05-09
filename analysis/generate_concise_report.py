"""
EMRI Concise Interim Report Generator (15-18 pages excluding Appendix)
Includes image references in Appendix section
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, bold=False, size=12):
    """Helper to set consistent font"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return run

def add_paragraph_with_style(doc, text, bold=False, first_indent=True, space_after=0):
    """Add paragraph with consistent styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p

def create_concise_report():
    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)

    # Configure headings
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.space_before = Pt(6)

    # ==================== TITLE PAGE ====================
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Data Analytics Capstone Topic")
    set_run_font(run, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Explainable Market Readiness Index (EMRI): A Threshold-Based Framework for Predicting High-Growth Sales Potential in Emerging Economies")
    set_run_font(run, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()

    # Interim Report
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Interim Report")
    set_run_font(run, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()

    # Author info
    info = ["Shankha Roy", "", "Walsh College", "QM640 V1: Data Analytics Capstone",
            "Dr. Vikas S", "Winter 2025 Term"]
    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run)

    doc.add_page_break()

    # ==================== GITHUB LINK ====================
    p = doc.add_paragraph()
    run = p.add_run("GitHub Repository")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc, "All data, code, and visualizations are available at: https://github.com/shankharoy/emri-capstone")
    doc.add_paragraph()

    # ==================== INTRODUCTION ====================
    doc.add_heading("Introduction", level=1)

    add_paragraph_with_style(doc,
        "Multinational FMCG firms deploy substantial capital toward emerging market expansion, yet prevailing assessment frameworks anchored in GDP per capita frequently misidentify optimal entry timing. "
        "This study develops the Explainable Market Readiness Index (EMRI), integrating digital infrastructure (internet penetration), urbanization rate, and female labor force participation into a threshold-based classification framework. "
        "The analysis covers 193 economies across 2018-2022, employing rigorous statistical methods to quantify structural readiness for e-commerce scalability.")

    # Scope and Objectives
    p = doc.add_paragraph()
    run = p.add_run("Scope and Objectives")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc, "This study aims to:")

    objectives = [
        "Quantify the relationship between digital infrastructure and e-commerce TAM using Spearman rank correlation;",
        "Assess workforce composition effects on consumer spending via Mann-Whitney U testing;",
        "Compare urbanization versus GNI per capita as market maturity predictors using OLS regression with Steiger's Z-test;",
        "Validate classification model performance achieving F1 ≥ 0.80 for high-growth market identification."
    ]

    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    # Research Problems
    p = doc.add_paragraph()
    run = p.add_run("Research Problems")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Four critical research problems inform the EMRI framework: (1) To what extent does digital adoption predict market potential? "
        "(2) Does female labor participation significantly differentiate consumer spending growth? "
        "(3) Which indicator—urbanization or GNI per capita—better predicts retail maturity? "
        "(4) Can statistical models classify high-growth markets with ≥0.80 accuracy?")

    doc.add_page_break()

    # ==================== LITERATURE SURVEY ====================
    doc.add_heading("Literature Survey", level=1)

    add_paragraph_with_style(doc,
        "The EMRI framework synthesizes four research streams. McKinsey Global Institute (2022) establishes digital infrastructure as foundational for retail ecosystems, "
        "with ITU (2023) demonstrating that 60-70% internet penetration thresholds trigger nonlinear e-commerce acceleration. "
        "World Bank (2024) validates female labor force participation as a predictor of household consumption, while OECD (2022) links rising FLFP to discretionary spending expansion. "
        "Urbanization literature confirms density's role in last-mile feasibility (World Bank, 2023). "
        "Methodologically, Green (1991) and Cohen (1988) inform power calculations and effect size interpretation, while Hosmer and Lemeshow (2000) guide logistic regression diagnostics. "
        "The explainable AI emphasis aligns with Breiman et al. (1984) and contemporary XAI scholarship prioritizing transparent decision rules for executive applications.")

    doc.add_page_break()

    # ==================== DATA DESCRIPTION ====================
    doc.add_heading("Data Description", level=1)

    add_paragraph_with_style(doc,
        "Data are sourced from World Bank World Development Indicators (WDI), covering 193 economies (2018-2022). "
        "The analytical sample comprises 1,275 country-year observations after KNN imputation (k=5). Five core indicators inform seven constructed EMRI indices.")

    # Table 1: Raw Variables
    p = doc.add_paragraph()
    run = p.add_run("Table 1")
    set_run_font(run, bold=True)
    p = doc.add_paragraph("World Bank WDI Indicators")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = doc.add_table(rows=6, cols=4)
    table1.style = 'Table Grid'

    headers = ['Variable', 'WDI Code', 'Description', 'Unit']
    for i, h in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    data = [
        ['gdp_per_capita', 'NY.GDP.PCAP.CD', 'GDP per capita, current USD', 'USD'],
        ['internet_access', 'IT.NET.USER.ZS', 'Internet users (% population)', '%'],
        ['urbanization', 'SP.URB.TOTL.IN.ZS', 'Urban population (% total)', '%'],
        ['female_labor', 'SL.TLF.CACT.FE.ZS', 'Female labor participation (%)', '%'],
        ['gni_per_capita', 'NY.GNP.PCAP.CD', 'GNI per capita, Atlas method', 'USD']
    ]

    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            table1.rows[i].cells[j].text = val
            for r in table1.rows[i].cells[j].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_paragraph()

    # Table 2: EMRI Indices
    p = doc.add_paragraph()
    run = p.add_run("Table 2")
    set_run_font(run, bold=True)
    p = doc.add_paragraph("EMRI Dimension Indices")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'

    for i, h in enumerate(['Index', 'Components', 'Weight', 'Description']):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    emri_data = [
        ['DII', 'Internet (70%), Urban (30%)', '25%', 'Digital infrastructure'],
        ['HCI', 'GDP (50%), Female Labor (50%)', '20%', 'Human capital'],
        ['ESI', 'GDP (50%), GNI (50%)', '20%', 'Economic stability'],
        ['GEI', 'Internet (40%), GDP (40%), Urban (20%)', '10%', 'Governance efficiency'],
        ['SRI', 'Urban (60%), Internet (40%)', '10%', 'Sustainability'],
        ['PSI', 'Log GNI per capita', '10%', 'Population scale'],
        ['SSI', 'GDP-GNI alignment', '5%', 'Security stability']
    ]

    for i, row in enumerate(emri_data, 1):
        for j, val in enumerate(row):
            table2.rows[i].cells[j].text = val
            for r in table2.rows[i].cells[j].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_page_break()

    # ==================== ANALYSIS ====================
    doc.add_heading("Analysis", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Data Cleaning and Preprocessing")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Raw data comprised 11,817 observations across 255 territories. Temporal filtering to 2018-2022 yielded 5,878 rows, pivoted to 1,275 country-year observations. "
        "Missing data ranged from 0% (urbanization) to 18.3% (internet users). KNN imputation (k=5, distance-weighted) recovered all missing values. "
        "Logarithmic transformation addressed extreme GDP/GNI skewness (skewness > 3.0).")

    p = doc.add_paragraph()
    run = p.add_run("Exploratory Data Analysis")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Distribution analysis revealed GDP per capita as heavily right-skewed (skewness = 3.28, kurtosis = 15.50), with mean ($17,717) nearly 3× median ($6,666), confirming extreme outlier influence. "
        "GNI per capita similarly exhibited extreme skew (2.86). In contrast, urbanization displayed near-normal distribution (skewness = -0.06), validating its use as a stable anchor variable. "
        "Internet penetration showed bimodal tendencies suggesting digital divide cohorts, while female labor participation clustered at 50-60% with lower-tail outliers indicating suppressed participation regions.")

    # Table 3: Distribution Stats
    p = doc.add_paragraph()
    run = p.add_run("Table 3")
    set_run_font(run, bold=True)
    p = doc.add_paragraph("Distribution Statistics Summary")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table3 = doc.add_table(rows=6, cols=5)
    table3.style = 'Table Grid'

    for i, h in enumerate(['Indicator', 'Mean', 'Median', 'Skewness', 'Kurtosis']):
        cell = table3.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    stats = [
        ['Internet Users (%)', '57.38', '63.22', '-0.32', '-1.18'],
        ['GDP per Capita', '$17,717', '$6,666', '3.28', '15.50'],
        ['GNI per Capita', '$15,201', '$6,199', '2.86', '4.37'],
        ['Female Labor (%)', '50.26', '52.15', '-0.66', '0.46'],
        ['Urbanization (%)', '60.45', '60.99', '-0.06', '-0.94']
    ]

    for i, row in enumerate(stats, 1):
        for j, val in enumerate(row):
            table3.rows[i].cells[j].text = val
            for r in table3.rows[i].cells[j].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_paragraph()

    add_paragraph_with_style(doc,
        "Correlation analysis revealed GDP-GNI near-collinearity (r = 0.985), necessitating exclusion of one variable. "
        "Spearman correlation between digital infrastructure and market readiness (ρ = 0.949) substantially exceeded Pearson (r = 0.893), confirming non-linear relationships better captured by rank-based methods. "
        "Female labor participation demonstrated near-zero correlation with economic indicators (|r| < 0.20), establishing orthogonality and unique contribution as a social readiness dimension.")

    doc.add_page_break()

    # ==================== MODELLING ====================
    doc.add_heading("Modelling", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Model Selection Rationale")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "The analytical pipeline employs explainable methods aligned with executive decision-support requirements. "
        "Spearman correlation addresses RQ1, accommodating non-normal distributions. Mann-Whitney U testing evaluates RQ2 without normality assumptions. "
        "OLS regression with Steiger's Z-test compares nested models for RQ3. Eight classification algorithms are evaluated for RQ4, with Logistic Regression selected for interpretability despite comparable performance across methods.")

    p = doc.add_paragraph()
    run = p.add_run("Feature Engineering")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Seven composite indices were constructed using min-max normalization (0-100 scale) and theoretically-derived weights. "
        "DII (25% weight) combines internet penetration and urbanization as the primary digital readiness indicator. "
        "HCI and ESI (20% each) capture human capital and economic stability. Four secondary indices (GEI, SRI, PSI, SSI) contribute 5-10% each. "
        "The binary target variable dichotomizes EMRI_Score at the median (High_Competitiveness: 1 = ≥50, 0 = <50).")

    p = doc.add_paragraph()
    run = p.add_run("Training and Validation")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Classification models employed 80/20 stratified train-test split (training: n=892; test: n=383), preserving class balance. "
        "Five-fold cross-validation provided robust performance estimates. Random state fixation (seed=42) ensures reproducibility. "
        "Continuous features were standardized prior to model fitting for distance-based and regularized algorithms.")

    doc.add_page_break()

    # ==================== PRELIMINARY RESULTS ====================
    doc.add_heading("Preliminary Results", level=1)

    # RQ1
    p = doc.add_paragraph()
    run = p.add_run("Research Question 1: Digital Infrastructure Correlation")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Spearman correlation between Digital Infrastructure Index (DII) and EMRI score yielded ρ = 0.949 (p < 0.001), exceeding the 0.60 threshold for practical significance. "
        "This very large effect size establishes digital infrastructure as the primary driver of market readiness. "
        "The Spearman coefficient substantially exceeded Pearson (r = 0.893), confirming that rank-based methods better capture the true non-linear relationship obscured by GDP outliers. "
        "H₀ (ρ = 0) is rejected; digital infrastructure is a necessary condition for high-growth potential.")

    # RQ2
    p = doc.add_paragraph()
    run = p.add_run("Research Question 2: Workforce Dynamics Comparison")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Mann-Whitney U testing revealed highly significant differences between developed and developing economies (U = 174,630, p < 0.001). "
        "Developed economies averaged EMRI = 79.2 versus 52.0 for developing economies, representing a 27.2-point gap. "
        "Effect size r = -0.814 indicates a large practical difference. H₀ (μ₁ = μ₂) is rejected; the digital divide is statistically and practically significant.")

    # Table 4
    p = doc.add_paragraph()
    run = p.add_run("Table 4")
    set_run_font(run, bold=True)
    p = doc.add_paragraph("Mann-Whitney U Test Results")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table4 = doc.add_table(rows=3, cols=4)
    table4.style = 'Table Grid'

    for i, h in enumerate(['Economy Status', 'Mean EMRI', 'Std Dev', 'N']):
        cell = table4.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    for i, row in enumerate([['Developed', '79.2', '8.4', '637'], ['Developing', '52.0', '12.7', '638']], 1):
        for j, val in enumerate(row):
            table4.rows[i].cells[j].text = val
            for r in table4.rows[i].cells[j].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_paragraph()

    # RQ3
    p = doc.add_paragraph()
    run = p.add_run("Research Question 3: Predictive Model Comparison")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "OLS regression compared simple (DII-only) versus full (7-index) models. Model 1 achieved R² = 0.893. Model 2 achieved R² = 1.000. "
        "Steiger's Z-test confirmed significant superiority of the full model (Z = inf, p < 0.001). H₀ (R²₁ = R²₂) is rejected; the multi-dimensional approach significantly improves prediction.")

    # Table 5
    p = doc.add_paragraph()
    run = p.add_run("Table 5")
    set_run_font(run, bold=True)
    p = doc.add_paragraph("OLS Model Comparison")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table5 = doc.add_table(rows=3, cols=5)
    table5.style = 'Table Grid'

    for i, h in enumerate(['Model', 'R-Squared', 'Adj R²', 'AIC', 'RMSE']):
        cell = table5.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    for i, row in enumerate([['Model 1 (Simple)', '0.893', '0.893', '8,391.5', '6.50'],
                              ['Model 2 (Full)', '1.000', '1.000', '-75,389.6', '0.00']], 1):
        for j, val in enumerate(row):
            table5.rows[i].cells[j].text = val
            for r in table5.rows[i].cells[j].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_paragraph()

    # RQ4
    p = doc.add_paragraph()
    run = p.add_run("Research Question 4: Classification Performance")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "Eight classification algorithms were evaluated. Logistic Regression achieved optimal performance: Accuracy = 0.997, Precision = 0.995, Recall = 1.000, F1 = 0.997, AUC-ROC = 1.000. "
        "Cross-validation confirmed stability (CV AUC = 1.000 ± 0.0001). The F1 = 0.997 substantially exceeds the 0.80 target threshold. "
        "All models significantly exceeded chance (p < 0.001). H₀ (AUC ≤ 0.50) is rejected; EMRI indices successfully classify market competitiveness.")

    # Table 6
    p = doc.add_paragraph()
    run = p.add_run("Table 6")
    set_run_font(run, bold=True)
    p = doc.add_paragraph("Classification Model Performance")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table6 = doc.add_table(rows=6, cols=4)
    table6.style = 'Table Grid'

    for i, h in enumerate(['Model', 'AUC-ROC', 'F1-Score', 'Accuracy']):
        cell = table6.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    results = [
        ['Logistic Regression', '1.000', '0.997', '0.997'],
        ['SVM (RBF)', '0.999', '0.987', '0.987'],
        ['KNN (k=5)', '0.999', '0.990', '0.990'],
        ['Gradient Boosting', '0.999', '0.982', '0.982'],
        ['Random Forest', '0.999', '0.984', '0.984']
    ]

    for i, row in enumerate(results, 1):
        for j, val in enumerate(row):
            table6.rows[i].cells[j].text = val
            for r in table6.rows[i].cells[j].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_paragraph()

    # Summary
    p = doc.add_paragraph()
    run = p.add_run("Statistical Hypothesis Summary")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "All four research hypotheses are rejected at α = 0.05. Digital infrastructure exhibits very strong correlation with market readiness (ρ = 0.949). "
        "Developed and developing economies differ significantly (27.2-point gap). The full 7-index model outperforms simple models. "
        "Classification achieves F1 = 0.997, exceeding the 0.80 target. The EMRI framework is statistically validated for strategic market prioritization.")

    doc.add_page_break()

    # ==================== BIBLIOGRAPHY ====================
    doc.add_heading("Bibliography", level=1)

    refs = [
        "Breiman, L., Friedman, J. H., Olshen, R. A., & Stone, C. J. (1984). Classification and regression trees. Wadsworth & Brooks/Cole.",
        "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.",
        "Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters, 27(8), 861–874.",
        "Green, S. B. (1991). How many subjects does it take to do a regression analysis? Multivariate Behavioral Research, 26(3), 499–510.",
        "Hosmer, D. W., & Lemeshow, S. (2000). Applied logistic regression (2nd ed.). John Wiley & Sons.",
        "International Telecommunication Union. (2023). Measuring digital development: Facts and figures 2023. ITU Publications.",
        "McKinsey Global Institute. (2022). The future of e-commerce in emerging markets. McKinsey & Company.",
        "OECD. (2022). Women at work in OECD countries: Progress and challenges. OECD Publishing.",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling with Python. Python in Science Conference.",
        "Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427–437.",
        "World Bank. (2023). World development indicators 2023. The World Bank Group. https://data.worldbank.org",
        "World Bank. (2024). Female labor force participation: Trends and drivers. World Bank Policy Research Working Paper."
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)

    doc.add_page_break()

    # ==================== APPENDIX ====================
    doc.add_heading("Appendix", level=1)

    p = doc.add_paragraph()
    run = p.add_run("A. Visualization Outputs and Figures")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "The following figures are generated from the analytical pipeline and available in the GitHub repository under /outputs/ and /notebooks/:")

    figures = [
        ("Figure A1", "Distribution Analysis by Indicator", "Histograms, KDE plots, box plots, and Q-Q plots showing normality, skewness, and outlier detection for all five WDI indicators (Internet Users, GDP per Capita, GNI per Capita, Female Labor Participation, Urbanization)."),
        ("Figure A2", "Pearson and Spearman Correlation Heatmaps", "Side-by-side correlation matrices comparing parametric (Pearson) and non-parametric (Spearman) correlations, revealing the non-linear relationship between GDP and Internet penetration (Spearman ρ = 0.868 vs Pearson r = 0.583)."),
        ("Figure A3", "Temporal Trends and Growth Trajectories", "Multi-panel time series showing GDP trajectories, Internet penetration trends, CAGR analysis, Female Labor participation, Urbanization patterns, and normalized global trends (2014-2023) for key economies (USA, CHN, IND, DEU, BRA, JPN)."),
        ("Figure A4", "Geographic Stratification by Income Tier", "Pie charts and box plots showing distribution by World Bank income classification, Digital Divide scatter plots (Internet vs GDP by tier), and Urbanization vs Female Labor participation patterns."),
        ("Figure A5", "RQ1: Spearman Correlation Analysis", "Scatter plot with trend line showing the relationship between Digital Infrastructure Index (DII) and EMRI Score, with correlation coefficient (ρ = 0.949) and significance annotation."),
        ("Figure A6", "RQ2: Mann-Whitney U Test Visualization", "Box plots comparing EMRI score distributions between Developed and Developing economies, showing median differences, IQR ranges, and outlier profiles."),
        ("Figure A7", "RQ3: OLS Regression Diagnostics", "Four-panel diagnostic plot including: (1) Residuals vs Fitted values, (2) Q-Q Plot for normality assessment, (3) Histogram of residuals, and (4) Actual vs Predicted scatter with R² annotation."),
        ("Figure A8", "RQ4: Classification Model Comparison", "Multi-panel visualization including: (1) ROC Curves for all 8 models, (2) Performance metrics bar charts, (3) Cross-validation AUC with error bars, and (4) Training Time vs Performance scatter plot."),
        ("Figure A9", "EMRI Dimension Index Distributions", "Distribution plots for all seven constructed indices (DII, HCI, ESI, GEI, SRI, PSI, SSI) showing min-max normalized ranges and density patterns."),
        ("Figure A10", "Market Priority Matrix", "Strategic quadrant analysis mapping economies by Digital Adoption and Economic Capacity, identifying 'Hidden Gems' and 'Saturated Giants' for investment prioritization.")
    ]

    for fig_num, title, desc in figures:
        p = doc.add_paragraph()
        run = p.add_run(f"{fig_num}: {title}")
        set_run_font(run, bold=True, size=11)

        p = doc.add_paragraph(desc)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)

    # Appendix B
    p = doc.add_paragraph()
    run = p.add_run("B. Model Coefficients and Decision Rules")
    set_run_font(run, bold=True, size=12)

    p = doc.add_paragraph()
    run = p.add_run("Logistic Regression Coefficients (for EMRI ≥ 50 classification):")
    set_run_font(run, bold=True, size=11)

    coeffs = [
        "Intercept: -0.156 (baseline odds)",
        "DII coefficient: 0.220 (OR = 1.246, p < 0.001)",
        "HCI coefficient: 0.200 (OR = 1.221, p < 0.001)",
        "ESI coefficient: 0.102 (OR = 1.107, p < 0.001)",
        "GEI coefficient: 0.157 (OR = 1.170, p < 0.001)",
        "SRI coefficient: 0.096 (OR = 1.101, p < 0.001)",
        "PSI coefficient: 0.123 (OR = 1.131, p < 0.001)",
        "SSI coefficient: 0.102 (OR = 1.107, p < 0.001)"
    ]

    for coeff in coeffs:
        p = doc.add_paragraph(coeff)
        p.paragraph_format.left_indent = Inches(0.5)
        for r in p.runs:
            r.font.name = 'Courier New'
            r.font.size = Pt(10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Example Decision Rule:")
    set_run_font(run, bold=True, size=11)

    p = doc.add_paragraph()
    run = p.add_run("IF DII > 68 AND HCI > 55 AND ESI > 45 THEN Classify as High Competitiveness (Probability > 0.85)")
    set_run_font(run, bold=False, size=10)
    for r in p.runs:
        r.font.name = 'Courier New'

    # Appendix C
    p = doc.add_paragraph()
    run = p.add_run("C. Data Processing Specifications")
    set_run_font(run, bold=True, size=12)

    specs = [
        "KNN Imputation: k=5 neighbors, Euclidean distance, distance-weighted averaging",
        "Log Transformation: Natural log applied to GDP and GNI per capita (ln(X))",
        "Min-Max Normalization: X_norm = (X - X_min) / (X_max - X_min) × 100",
        "Outlier Treatment: IQR method (1.5×IQR fence) for outlier identification",
        "Train-Test Split: 80/20 stratified split with random_state=42",
        "Cross-Validation: 5-fold stratified CV for model evaluation",
        "Significance Level: α = 0.05 for all hypothesis tests",
        "Effect Size: Cohen's conventions (r = 0.1 small, 0.3 medium, ≥0.5 large)"
    ]

    for spec in specs:
        p = doc.add_paragraph(spec, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    # Appendix D
    p = doc.add_paragraph()
    run = p.add_run("D. Repository Structure and File Locations")
    set_run_font(run, bold=True, size=12)

    add_paragraph_with_style(doc,
        "All data files, code notebooks, and visualizations are organized in the GitHub repository (https://github.com/shankharoy/emri-capstone) as follows:")

    structure = [
        "data/output/wdi_panel.csv — Raw World Bank WDI panel data (11,817 observations)",
        "notebooks/wdi_eda_professional.ipynb — Comprehensive EDA with visualizations",
        "notebooks/EMAR_Complete_Analytical_Pipeline.ipynb — Full modeling pipeline",
        "outputs/correlation_heatmaps.png — Pearson and Spearman correlation matrices",
        "outputs/emri_distributions.png — Distribution analysis of EMRI indices",
        "outputs/rq1_spearman_analysis.png — RQ1 correlation visualization",
        "outputs/rq2_mann_whitney.png — RQ2 group comparison plots",
        "outputs/rq3_ols_regression.png — RQ3 regression diagnostics",
        "outputs/rq4_classification_comparison.png — RQ4 model performance comparison"
    ]

    for item in structure:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        for r in p.runs:
            r.font.name = 'Courier New'
            r.font.size = Pt(10)

    return doc

# Generate the report
print("Generating concise 15-18 page Interim Report...")
doc = create_concise_report()

# Save
output_path = "D:\\shankha\\github\\emri-capstone\\QM640_Interim_Report_Shankha_Roy_Concise.docx"
doc.save(output_path)
print(f"Concise Interim Report saved to: {output_path}")
