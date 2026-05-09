"""
EMRI Interim Report Generator
Generates a comprehensive, publication-quality Interim Report following QM640 template requirements
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_border(cell, **kwargs):
    """Set cell border properties"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcPr.append(edge_el)

def create_interim_report():
    # Create document
    doc = Document()

    # Set default font and spacing for the document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)

    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        heading_style.paragraph_format.space_after = Pt(12)
        heading_style.paragraph_format.space_before = Pt(12)

    # ==================== TITLE PAGE ====================
    # Add spacing for vertical centering
    for _ in range(8):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Data Analytics Capstone Topic")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Explainable Market Readiness Index (EMRI): A Threshold-Based Framework for Predicting High-Growth Sales Potential in Emerging Economies")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Interim Report label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Interim Report")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Author information
    info_lines = [
        "Shankha Roy",
        "",
        "Walsh College",
        "QM640 V1: Data Analytics Capstone",
        "Dr. Vikas S",
        "Winter 2025 Term"
    ]

    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # Page break after title page
    doc.add_page_break()

    # ==================== GITHUB LINK SECTION ====================
    p = doc.add_paragraph()
    run = p.add_run("GitHub Repository and Data Access")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    p = doc.add_paragraph()
    run = p.add_run("All data files, code notebooks, and supplementary materials are available at the following repository: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run("https://github.com/shankharoy/emri-capstone")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True

    p = doc.add_paragraph()
    run = p.add_run("The repository contains the complete analytical pipeline, including World Bank WDI extraction scripts, KNN imputation modules, EMRI index construction algorithms, and statistical modeling notebooks. All data are sourced from publicly available World Bank Development Indicators (2018-2022), ensuring full reproducibility and transparency.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_page_break()

    # ==================== INTRODUCTION SECTION ====================
    heading = doc.add_heading("Introduction", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Background and Context")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    intro_text = """Multinational fast-moving consumer goods (FMCG) firms collectively deploy substantial capital toward market entry and expansion initiatives, with significant investments concentrated in emerging economies across Southeast Asia, Sub-Saharan Africa, and South Asia (World Bank, 2023). Despite the magnitude of these commitments, prevailing market assessment frameworks remain anchored in lagging macroeconomic indicators, most notably Gross Domestic Product (GDP) per capita, which reflect historical economic performance rather than forward-looking structural readiness. This reliance on retrospective indicators has produced a persistent pattern of strategic misalignment: firms frequently enter markets that appear economically attractive in aggregate but lack the digital infrastructure, urban density, and consumer base composition required to support scalable retail operations within the investment horizon (McKinsey Global Institute, 2022).

The accelerating digitalization of consumer commerce has fundamentally reshaped the prerequisites for successful market entry. Internet penetration, mobile connectivity, and urban logistical density now function as enabling conditions for e-commerce viability in ways that GDP alone cannot anticipate. Evidence from the International Telecommunication Union (ITU, 2023) demonstrates that economies surpassing the 60-70% internet penetration threshold experience nonlinear increases in e-commerce transaction volume, an inflection point that GDP-based metrics fail to predict with the requisite two-to-three-year lead time for effective capital allocation. Consequently, FMCG firms face a systematic market-timing challenge: investments are often deployed prematurely, before digital infrastructure reaches commercial viability, or belatedly, after competitors have established distribution networks and brand equity.

Beyond digital infrastructure, the socioeconomic composition of the consumer base introduces additional structural determinants that aggregate GDP measures obscure. Female labour force participation (FLFP) serves as a robust leading indicator of dual-income household formation and the associated expansion in discretionary consumption (World Bank, 2024). Rising FLFP trajectories signal the emergence of consumer segments that allocate a larger share of incremental income to branded FMCG categories, including packaged foods, personal care, and household goods (OECD, 2022). Similarly, urbanization rates shape the economic feasibility of last-mile fulfillment and retail distribution, particularly in countries where rural populations exceed half of the total population. These structural factors collectively determine whether a market can support profitable, scalable FMCG e-commerce operations, yet they remain underrepresented in traditional market readiness assessments."""

    p = doc.add_paragraph(intro_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Scope and Objectives
    p = doc.add_paragraph()
    run = p.add_run("Scope and Objectives")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    scope_text = """The purpose of this study is to develop and validate the Explainable Market Readiness Index (EMRI), a quantitative framework for classifying emerging economies based on their structural readiness to support profitable FMCG e-commerce operations. The study analyzes 193 economies across the 2018-2022 longitudinal panel, using World Bank Development Indicators as primary variables. The scope is intentionally limited to structural predictors of market readiness that are publicly verifiable, annually updated, and available for the full cross-section of World Bank member economies.

The specific objectives of this research are to:"""

    p = doc.add_paragraph(scope_text)

    objectives = [
        "Quantify the relationship between digital infrastructure adoption and the Total Addressable Market (TAM) for e-commerce, operationalized through Spearman rank correlation analysis;",
        "Assess whether workforce composition, measured through female labour force participation, produces statistically significant differences in consumer spending growth, evaluated via Mann-Whitney U testing;",
        "Determine whether urbanization rate or Gross National Income (GNI) per capita serves as the stronger predictor of retail market maturity, employing Ordinary Least Squares (OLS) regression with Steiger's Z-test for dependent correlation comparison;",
        "Evaluate whether a transparent, statistically grounded classification model can accurately identify high-growth FMCG e-commerce markets, achieving an F1-score of at least 0.80 using only publicly available structural indicators."
    ]

    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    # Research Problems
    p = doc.add_paragraph()
    run = p.add_run("Research Problems Identified")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    problem_text = """This study addresses four critical research problems that collectively inform the EMRI framework:

Research Problem 1: Digital Infrastructure and Market Potential. To what extent does the Digital Adoption Index (Internet/Mobile penetration) predict the total addressable market for e-commerce sales in emerging economies? This problem examines whether digital infrastructure serves as a necessary condition for market entry viability.

Research Problem 2: Workforce Dynamics and Consumer Spending. Is there a statistically significant difference in consumer spending growth between nations with high versus low female labour force participation? This investigation validates workforce composition as a secondary EMRI filter for dual-income household effects.

Research Problem 3: Urban Density as Market Maturity Signal. Which specific socioeconomic indicator, urbanization rate or GNI per capita, serves as the stronger predictor of retail market maturity? This comparative analysis establishes the dominant structural driver of logistical scalability.

Research Problem 4: Predictive Classification Accuracy. Can a statistical or machine learning model classify high-growth markets with accuracy exceeding 0.80? This validation problem confirms the EMRI's utility as a decision-support tool for strategic market prioritization."""

    p = doc.add_paragraph(problem_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ==================== LITERATURE SURVEY ====================
    heading = doc.add_heading("Literature Survey", level=1)

    lit_text = """The development of the Explainable Market Readiness Index (EMRI) draws upon four interconnected streams of academic and practitioner literature: digital economy frameworks, gender economics and workforce participation, urbanization and infrastructure development, and explainable machine learning for strategic decision-making.

Digital Economy and Market Readiness. McKinsey Global Institute (2022) establishes that digital infrastructure serves as the foundational enabler for modern retail ecosystems, with internet penetration functioning as a threshold variable rather than a linear predictor of e-commerce viability. The International Telecommunication Union (ITU, 2023) demonstrates that economies exceeding 60-70% internet penetration experience nonlinear acceleration in digital transaction volumes, creating distinct "digital readiness" categories that traditional GDP-centric models fail to capture. This literature supports the EMRI's emphasis on structural digital indicators as leading predictors of market potential.

Gender Economics and Consumer Behavior. World Bank (2024) research establishes female labour force participation (FLFP) as a robust predictor of household consumption patterns, with dual-income households demonstrating significantly higher discretionary spending on branded consumer goods. OECD (2022) findings indicate that rising FLFP correlates with expanded market opportunities for FMCG categories, particularly in emerging economies undergoing demographic transitions. This body of work justifies the inclusion of FLFP as an independent dimension within the EMRI framework, capturing social readiness that GDP metrics obscure.

Urbanization and Logistical Infrastructure. The relationship between urban density and retail market maturity has been extensively documented in development economics literature. World Bank Development Indicators (2023) demonstrate that urbanization rates above 50% correlate with improved last-mile delivery feasibility and reduced logistical costs for e-commerce operations. This research stream informs the EMRI's urbanization component, which serves as a proxy for physical infrastructure readiness complementary to digital connectivity measures.

Explainable AI in Strategic Decision-Making. The methodological emphasis on interpretable statistical models aligns with emerging literature on explainable artificial intelligence (XAI) for business applications. Breiman et al. (1984) foundational work on classification and regression trees established the importance of transparent, rule-based models for executive decision-making. Recent scholarship emphasizes that depth-constrained decision trees and logistic regression with odds ratios provide actionable intelligence superior to black-box ensemble methods for strategic market entry decisions (Pedregosa et al., 2011; Seabold & Perktold, 2010).

Statistical Methodology in Development Economics. Green (1991) sample size guidelines inform the study's power calculations, ensuring adequate statistical power across all four research questions. Cohen (1988) effect size frameworks guide the interpretation of Spearman correlations and Mann-Whitney U tests, establishing thresholds for practical significance beyond mere statistical significance. Hosmer and Lemeshow (2000) logistic regression diagnostics inform the classification model evaluation, while Fawcett (2006) ROC analysis principles guide AUC interpretation for binary classification tasks.

Market Entry Strategy. Contemporary strategic management literature emphasizes the importance of forward-looking structural indicators in emerging market entry decisions. The EMRI responds to identified gaps in existing frameworks by integrating three distinct readiness dimensions: digital infrastructure (velocity), urbanization (structural stability), and workforce composition (social readiness). This multi-dimensional approach addresses limitations identified in GDP-centric assessment tools while maintaining the interpretability required for board-level strategic planning."""

    p = doc.add_paragraph(lit_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ==================== DATA DESCRIPTION ====================
    heading = doc.add_heading("Data Description", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Data Sources")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    data_source_text = """All data are sourced exclusively from publicly accessible, institutionally verified repositories. The primary data source is the World Bank Open Data portal, specifically the World Development Indicators (WDI) database, accessible at https://data.worldbank.org/indicator. The dataset comprises five core indicators extracted for 193 economies across the 2018-2022 temporal window, yielding 1,275 country-year observations following data quality procedures."""

    p = doc.add_paragraph(data_source_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Create data dictionary table
    p = doc.add_paragraph()
    run = p.add_run("Table 1")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    p = doc.add_paragraph("World Bank WDI Indicators: Data Dictionary")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create table for raw variables
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ['Variable Name', 'WDI Code', 'Description', 'Type/Unit']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    # Data rows
    data_rows = [
        ['gdp_per_capita', 'NY.GDP.PCAP.CD', 'GDP per capita, current USD. Primary proxy for e-commerce TAM and market maturity.', 'Continuous/USD'],
        ['internet_access', 'IT.NET.USER.ZS', 'Individuals using the Internet as % of total population.', 'Continuous/%'],
        ['urbanization', 'SP.URB.TOTL.IN.ZS', 'Urban population as % of total population. Proxy for logistical density.', 'Continuous/%'],
        ['female_labor', 'SL.TLF.CACT.FE.ZS', 'Female labor force participation rate, ages 15+, as % of female population.', 'Continuous/%'],
        ['gni_per_capita', 'NY.GNP.PCAP.CD', 'GNI per capita, Atlas method, USD. Used for comparative analysis.', 'Continuous/USD']
    ]

    for i, row_data in enumerate(data_rows, 1):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()  # Space after table

    p = doc.add_paragraph()
    run = p.add_run("Constructed Variables: EMRI Dimension Indices")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    emri_text = """Following data extraction, seven composite indices were constructed using min-max normalization to 0-100 scales, combined using theoretically-derived weights. These indices form the analytical foundation for the EMRI framework:"""

    p = doc.add_paragraph(emri_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # EMRI indices table
    p = doc.add_paragraph()
    run = p.add_run("Table 2")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    p = doc.add_paragraph("EMRI Dimension Indices: Construction and Weights")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'

    headers2 = ['Index', 'Components', 'Weight', 'Description']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    emri_rows = [
        ['DII (Digital Infrastructure)', 'Internet (70%), Urban (30%)', '25%', 'Core digital infrastructure readiness'],
        ['HCI (Human Capital)', 'GDP (50%), Female Labor (50%)', '20%', 'Development and inclusion proxy'],
        ['ESI (Economic Stability)', 'GDP (50%), GNI (50%)', '20%', 'Output-income consistency measure'],
        ['GEI (Governance Efficiency)', 'Internet (40%), GDP (40%), Urban (20%)', '10%', 'Digital-wealth connectivity proxy'],
        ['SRI (Sustainability)', 'Urban (60%), Internet (40%)', '10%', 'Environmental efficiency readiness'],
        ['PSI (Population Scale)', 'Log GNI per capita', '10%', 'Market size indicator'],
        ['SSI (Security Stability)', 'GDP-GNI alignment', '5%', 'Economic diversity proxy']
    ]

    for i, row_data in enumerate(emri_rows, 1):
        for j, cell_text in enumerate(row_data):
            cell = table2.rows[i].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # Target variable
    p = doc.add_paragraph()
    run = p.add_run("Target Variable")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    target_text = """The binary target variable, High_Competitiveness, was constructed by dichotomizing EMRI_Score at the median threshold. Observations scoring at or above the median receive a value of 1 (High-Growth Potential), while those below receive 0 (Low-Growth Potential). This approach ensures balanced class representation (approximately 50/50 split) suitable for classification model training and evaluation."""

    p = doc.add_paragraph(target_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # Continue with Analysis section...
    print("Document generation in progress... Part 1 completed")

    return doc

def add_analysis_section(doc):
    """Add Analysis section with EDA findings"""

    heading = doc.add_heading("Analysis", level=1)

    # Data Cleaning
    p = doc.add_paragraph()
    run = p.add_run("Data Cleaning and Preprocessing")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    cleaning_text = """The raw dataset contained 11,817 observations across 255 countries/territories. Following extraction and temporal filtering to the 2018-2022 analysis window, the dataset comprised 5,878 rows in long format. Pivot operations transformed these data into wide format with 1,275 country-year observations across eight variables.

Missing data analysis revealed differential completeness patterns across indicators. Urban population achieved 100% completeness (1,275 observations), while internet users exhibited 18.3% missingness (233 observations), and female labor participation demonstrated 11.8% missingness (151 observations). Rather than employing listwise deletion, which would reduce analytical power, missing values were imputed using K-Nearest Neighbors (KNN) with k=5 and distance-weighted averaging. This approach leverages economically similar regional neighbors to estimate missing values while preserving the underlying data structure. Following imputation, the analytical sample contained 1,275 complete observations with 100% data coverage."""

    p = doc.add_paragraph(cleaning_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # EDA Results
    p = doc.add_paragraph()
    run = p.add_run("Exploratory Data Analysis Results")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    eda_text = """Distribution analysis revealed significant non-normality in economic indicators, necessitating methodological adjustments for downstream modeling. GDP per capita exhibited extreme positive skewness (3.27) with mean ($17,716) nearly three times the median ($6,665), indicating a "winner-takes-all" economic structure where high-income outliers distort central tendency measures. GNI per capita similarly demonstrated extreme right skew (2.86), confirming that monetary metrics require logarithmic transformation before parametric modeling.

In contrast, structural indicators displayed more favorable distributional properties. Urbanization rate exhibited near-normal distribution (skewness = -0.05), validating its use as a stable "anchor" variable for cross-country comparisons. Internet penetration demonstrated modest negative skewness (-0.32) with a bimodal tendency suggesting a global digital divide between "Highly Connected" and "Developing Connectivity" cohorts. Female labor participation showed moderate negative skewness (-0.65), with clustering around 50-60% participation rates and lower-tail outliers indicating regions with structurally suppressed female economic engagement."""

    p = doc.add_paragraph(eda_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Create EDA results table
    p = doc.add_paragraph()
    run = p.add_run("Table 3")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    p = doc.add_paragraph("Distribution Statistics Summary by Indicator")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table3 = doc.add_table(rows=6, cols=6)
    table3.style = 'Table Grid'

    headers3 = ['Indicator', 'Mean', 'Median', 'Std Dev', 'Skewness', 'Kurtosis']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    stats_rows = [
        ['Internet Users (%)', '57.38', '63.22', '28.36', '-0.322', '-1.180'],
        ['GDP per Capita (US$)', '17,716.82', '6,665.74', '26,664.21', '3.275', '15.498'],
        ['GNI per Capita (US$)', '15,200.77', '6,199.06', '20,346.01', '2.061', '4.365'],
        ['Female Labor Participation (%)', '50.26', '52.15', '14.94', '-0.658', '0.459'],
        ['Urban Population (%)', '60.45', '60.99', '22.61', '-0.056', '-0.939']
    ]

    for i, row_data in enumerate(stats_rows, 1):
        for j, cell_text in enumerate(row_data):
            cell = table3.rows[i].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # Correlation Analysis
    p = doc.add_paragraph()
    run = p.add_run("Correlation Structure and Multicollinearity Assessment")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    corr_text = """Bivariate correlation analysis revealed both expected relationships and potential modeling concerns. Pearson correlation between GDP and GNI per capita reached 0.985, indicating virtual identity between these measures and necessitating the exclusion of one variable from predictive models to prevent multicollinearity inflation of standard errors.

The relationship between digital infrastructure and economic output demonstrated methodological sensitivity. Pearson correlation between internet penetration and GDP per capita registered 0.583, while Spearman rank correlation for the same relationship reached 0.868. This substantial divergence (Δρ = 0.285) confirms that the relationship is strongly monotonic but non-linear, with Spearman's rank-based approach better capturing the true association strength obscured by GDP's extreme outliers. Female labor participation exhibited near-zero correlation with all other indicators (|r| < 0.20), establishing its orthogonality and unique contribution as a "Social Readiness" dimension independent of wealth effects."""

    p = doc.add_paragraph(corr_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Temporal Analysis
    p = doc.add_paragraph()
    run = p.add_run("Temporal Trends and Growth Trajectories")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    temporal_text = """Longitudinal analysis across the 2014-2023 window revealed divergent growth patterns across indicators. Compound Annual Growth Rate (CAGR) calculations identified China and India as high-momentum economies with growth rates approaching 6%, while Nigeria, Brazil, and Japan exhibited negative or stagnant trajectories. Internet penetration demonstrated the most aggressive growth, with India exemplifying "digital leapfrogging" by increasing connectivity from approximately 15% to 60% within the decade. This velocity significantly exceeds GDP growth rates, confirming digital infrastructure as a leading rather than lagging indicator of market readiness.

Normalized trend analysis (2014 = 100) positioned internet users as the dominant growth indicator (60% increase), while urbanization and female labor participation remained essentially flat (0% net change). GDP and GNI exhibited significant V-shaped volatility around 2020, reflecting COVID-19 economic shocks followed by sharp recovery. These patterns inform the EMRI weighting framework: velocity metrics (internet growth, GDP CAGR) receive higher weights for short-term prediction, while structural metrics (urbanization) serve as baseline filters for market entry viability."""

    p = doc.add_paragraph(temporal_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Geographic Stratification
    p = doc.add_paragraph()
    run = p.add_run("Geographic Stratification by Income Tier")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True

    geo_text = """World Bank income tier classification revealed clear structural stratification. High-income economies (n=86) averaged $43,111 GDP per capita with 84.2% internet penetration and 77.2% urbanization. Low-income economies (n=26) averaged $741 GDP with 16.0% internet penetration and 34.8% urbanization. The middle-income tiers demonstrated the most dynamic characteristics: upper-middle income economies (n=74) showed 62.1% internet penetration with substantial variance, suggesting transitional digital adoption patterns ideal for EMRI targeting.

Digital divide analysis confirmed an S-curve relationship between income and connectivity. Internet growth proves most explosive in middle-income tiers, while high-income markets exhibit saturation (80-100% penetration). This non-linearity validates the EMRI's focus on middle-income emerging markets where digital infrastructure offers maximum discriminatory power for investment decisions."""

    p = doc.add_paragraph(geo_text)
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    return doc

# Create the document
doc = create_interim_report()
doc = add_analysis_section(doc)

# Save the document
output_path = "D:\\shankha\\github\\emri-capstone\\QM640_Interim_Report_Shankha_Roy.docx"
doc.save(output_path)
print(f"Interim Report (Part 1) saved to: {output_path}")
